"""
Appeal Document Generator for FasalPramaan
Generates professional PDF and Word documents for insurance claim appeals
"""

from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import io
import logging

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class AppealGenerator:
    """Generates professional appeal documents for crop insurance claims"""
    
    def __init__(self):
        self.styles = getSampleStyleSheet()
        self._setup_custom_styles()
    
    def _setup_custom_styles(self):
        """Create custom styles for PDF generation"""
        # Title style
        self.title_style = ParagraphStyle(
            'CustomTitle',
            parent=self.styles['Heading1'],
            fontSize=18,
            textColor=colors.darkgreen,
            alignment=TA_CENTER,
            spaceAfter=12
        )
        
        # Heading style
        self.heading_style = ParagraphStyle(
            'CustomHeading',
            parent=self.styles['Heading2'],
            fontSize=14,
            textColor=colors.darkblue,
            spaceAfter=6
        )
        
        # Sub-heading style
        self.subheading_style = ParagraphStyle(
            'CustomSubheading',
            parent=self.styles['Heading3'],
            fontSize=12,
            textColor=colors.darkgrey,
            spaceAfter=4
        )
        
        # Body style
        self.body_style = ParagraphStyle(
            'CustomBody',
            parent=self.styles['Normal'],
            fontSize=10,
            leading=14,
            spaceAfter=8
        )
        
        # Bullet style
        self.bullet_style = ParagraphStyle(
            'CustomBullet',
            parent=self.styles['Normal'],
            fontSize=10,
            leading=14,
            spaceAfter=4,
            leftIndent=20
        )
        
        # Signature style
        self.signature_style = ParagraphStyle(
            'CustomSignature',
            parent=self.styles['Normal'],
            fontSize=10,
            leading=14,
            spaceAfter=4,
            alignment=TA_RIGHT
        )
    
    def generate_pdf(self, data: dict) -> bytes:
        """Generate a professional PDF appeal document"""
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=72
        )
        
        story = []
        
        # Add header
        story.append(Paragraph("FASALPRAMAAN", self.title_style))
        story.append(Paragraph("Independent Crop Insurance Verification", self.subheading_style))
        story.append(Spacer(1, 0.2 * inch))
        story.append(Paragraph("=" * 80, self.body_style))
        story.append(Spacer(1, 0.3 * inch))
        
        # Title: Appeal Letter
        story.append(Paragraph("APPEAL LETTER", self.heading_style))
        story.append(Spacer(1, 0.1 * inch))
        
        # Date
        story.append(Paragraph(f"Date: {datetime.now().strftime('%B %d, %Y')}", self.body_style))
        story.append(Spacer(1, 0.2 * inch))
        
        # Recipient
        story.append(Paragraph("To:", self.heading_style))
        story.append(Paragraph("The District Grievance Redressal Committee", self.body_style))
        story.append(Paragraph("[District Name], [State]", self.body_style))
        story.append(Spacer(1, 0.1 * inch))
        
        # Subject
        story.append(Paragraph(f"Subject: Appeal against crop insurance assessment for {data.get('plot_name', 'Plot')}", self.heading_style))
        story.append(Spacer(1, 0.2 * inch))
        
        # Body: Introduction
        story.append(Paragraph("Respected Members of the Committee,", self.body_style))
        story.append(Spacer(1, 0.1 * inch))
        
        intro = f"""I am writing to formally appeal the crop insurance assessment for my plot located at 
latitude {data.get('latitude', 'N/A')}, longitude {data.get('longitude', 'N/A')} for the {data.get('season', 'Kharif 2017')} season. 
Based on independent analysis using satellite imagery and weather data, I am submitting the following evidence 
for your review."""
        story.append(Paragraph(intro, self.body_style))
        story.append(Spacer(1, 0.2 * inch))
        
        # Section 1: Satellite Analysis
        story.append(Paragraph("1. SATELLITE-BASED VEGETATION ANALYSIS", self.heading_style))
        story.append(Paragraph(f"""Using Sentinel-2 satellite imagery accessed through Google Earth Engine, 
the Normalized Difference Vegetation Index (NDVI) was calculated for the period of {data.get('damage_period', 'damage period')}. 
NDVI is a widely accepted, peer-reviewed measure of vegetation health.""", self.body_style))
        story.append(Spacer(1, 0.05 * inch))
        
        # Deviation score in a table
        deviation = data.get('deviation_score', 0)
        if deviation < -2:
            status_text = 'Severe'
        elif deviation < -1.5:
            status_text = 'Significant'
        else:
            status_text = 'Normal'
        
        data_table = [
            ['Metric', 'Value', 'Status'],
            ['NDVI Deviation Score', f"{deviation:.2f} \u03C3", status_text],
            ['Cloud-free Images Used', str(data.get('image_count', 0)), '\u2713'],
            ['Average Cloud Cover', f"{data.get('cloud_cover_avg', 0):.1f}%", '\u2713']
        ]
        
        table = Table(data_table, colWidths=[2.5 * inch, 1.5 * inch, 1.5 * inch])
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.darkgreen),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 10),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        story.append(table)
        story.append(Spacer(1, 0.2 * inch))
        
        # Section 2: Weather Analysis
        story.append(Paragraph("2. WEATHER DATA ANALYSIS", self.heading_style))
        story.append(Paragraph("Weather data was obtained from the Open-Meteo API using ERA5-Land reanalysis data.", self.body_style))
        story.append(Spacer(1, 0.05 * inch))
        
        # Weather data as bullet points - ONE bullet per paragraph
        story.append(Paragraph(f"\u2022 Total Rainfall: {data.get('rainfall_total', 0):.1f} mm", self.bullet_style))
        story.append(Paragraph(f"\u2022 Rainy Days: {data.get('rainfall_days', 0)} days", self.bullet_style))
        story.append(Paragraph(f"\u2022 Comparison: {data.get('rainfall_comparison', 'No data')}", self.bullet_style))
        story.append(Paragraph(f"\u2022 Data Source: {data.get('weather_source', 'ERA5-Land')}", self.bullet_style))
        story.append(Spacer(1, 0.05 * inch))
        story.append(Paragraph("This weather analysis provides independent verification of the satellite observations.", self.body_style))
        story.append(Spacer(1, 0.2 * inch))
        
        # Section 3: Conclusion
        story.append(Paragraph("3. CONCLUSION", self.heading_style))
        
        # Get rainfall comparison without duplication
        rainfall_text = data.get('rainfall_comparison', '')
        if "This independent evidence" in rainfall_text:
            rainfall_text = rainfall_text.split("This independent evidence")[0].strip()
        
        deviation = data.get('deviation_score', 0)
        if data.get('status') == 'anomaly_detected':
            severity = 'severe' if deviation < -2 else 'significant'
            conclusion = f"""The analysis shows a significant deviation from normal vegetation patterns. 
The NDVI deviation score of {deviation:.2f} standard deviations below the historical baseline 
indicates {severity} vegetation stress.
{rainfall_text}
This independent evidence supports the claim of crop damage and warrants a re-evaluation of the insurance assessment."""
        else:
            conclusion = f"""The analysis shows vegetation health within normal range for this season.
The NDVI deviation score of {deviation:.2f} standard deviations from the historical baseline
indicates normal vegetation conditions.
{rainfall_text}
This independent evidence suggests the insurance assessment should be reviewed in light of normal vegetation conditions."""
        
        # Replace newlines with <br/> for PDF
        conclusion = conclusion.replace('\n', '<br/>')
        story.append(Paragraph(conclusion, self.body_style))
        story.append(Spacer(1, 0.2 * inch))
        
        # Section 4: Request
        story.append(Paragraph("4. REQUEST", self.heading_style))
        request_text = """In light of the evidence presented above, I respectfully request the following:<br/><br/>
1. A formal review of the crop insurance claim assessment for the plot described above<br/>
2. Consideration of the independent satellite and weather data as evidence<br/>
3. A fair and transparent evaluation based on all available evidence<br/><br/>
I am available to provide any additional information or clarification as needed."""
        story.append(Paragraph(request_text, self.body_style))
        story.append(Spacer(1, 0.2 * inch))
        
        # Signature
        story.append(Spacer(1, 0.5 * inch))
        story.append(Paragraph("Sincerely,", self.body_style))
        story.append(Spacer(1, 0.1 * inch))
        story.append(Paragraph("[Farmer Name]", self.body_style))
        story.append(Paragraph("[Contact Number]", self.body_style))
        story.append(Paragraph("[Email Address]", self.body_style))
        story.append(Paragraph("[Village], [District], [State]", self.body_style))
        
        # Build PDF
        doc.build(story)
        buffer.seek(0)
        return buffer.getvalue()
    
    def generate_word(self, data: dict) -> bytes:
        """Generate a professional Word document appeal"""
        doc = Document()
        
        # Title
        title = doc.add_heading('FASALPRAMAAN', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_heading('Independent Crop Insurance Verification', level=1)
        doc.add_paragraph('=' * 60)
        
        # Date
        doc.add_paragraph(f"Date: {datetime.now().strftime('%B %d, %Y')}")
        doc.add_paragraph()
        
        # To
        doc.add_heading('To:', level=2)
        doc.add_paragraph('The District Grievance Redressal Committee')
        doc.add_paragraph('[District Name], [State]')
        doc.add_paragraph()
        
        # Subject
        doc.add_heading(f"Subject: Appeal against crop insurance assessment for {data.get('plot_name', 'Plot')}", level=2)
        doc.add_paragraph()
        
        # Introduction
        doc.add_paragraph('Respected Members of the Committee,')
        doc.add_paragraph()
        
        intro = f"""I am writing to formally appeal the crop insurance assessment for my plot located at 
latitude {data.get('latitude', 'N/A')}, longitude {data.get('longitude', 'N/A')} for the {data.get('season', 'Kharif 2017')} season. 
Based on independent analysis using satellite imagery and weather data, I am submitting the following evidence 
for your review."""
        doc.add_paragraph(intro)
        doc.add_paragraph()
        
        # Satellite Analysis
        doc.add_heading('1. SATELLITE-BASED VEGETATION ANALYSIS', level=2)
        doc.add_paragraph(f"""Using Sentinel-2 satellite imagery accessed through Google Earth Engine, 
the Normalized Difference Vegetation Index (NDVI) was calculated for the period of {data.get('damage_period', 'damage period')}. 
NDVI is a widely accepted, peer-reviewed measure of vegetation health.""")
        doc.add_paragraph()
        
        # Add table
        table = doc.add_table(rows=4, cols=3)
        table.style = 'Table Grid'
        
        headers = ['Metric', 'Value', 'Status']
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            cell.paragraphs[0].runs[0].bold = True
        
        deviation = data.get('deviation_score', 0)
        if deviation < -2:
            status_text = 'Severe'
        elif deviation < -1.5:
            status_text = 'Significant'
        else:
            status_text = 'Normal'
        
        table.rows[1].cells[0].text = 'NDVI Deviation Score'
        table.rows[1].cells[1].text = f"{deviation:.2f} σ"
        table.rows[1].cells[2].text = status_text
        
        table.rows[2].cells[0].text = 'Cloud-free Images Used'
        table.rows[2].cells[1].text = str(data.get('image_count', 0))
        table.rows[2].cells[2].text = '✓'
        
        table.rows[3].cells[0].text = 'Average Cloud Cover'
        table.rows[3].cells[1].text = f"{data.get('cloud_cover_avg', 0):.1f}%"
        table.rows[3].cells[2].text = '✓'
        
        doc.add_paragraph()
        
        # Weather Analysis
        doc.add_heading('2. WEATHER DATA ANALYSIS', level=2)
        doc.add_paragraph("Weather data was obtained from the Open-Meteo API using ERA5-Land reanalysis data.")
        doc.add_paragraph(f"• Total Rainfall: {data.get('rainfall_total', 0):.1f} mm")
        doc.add_paragraph(f"• Rainy Days: {data.get('rainfall_days', 0)} days")
        doc.add_paragraph(f"• Comparison: {data.get('rainfall_comparison', 'No data')}")
        doc.add_paragraph(f"• Data Source: {data.get('weather_source', 'ERA5-Land')}")
        doc.add_paragraph()
        doc.add_paragraph("This weather analysis provides independent verification of the satellite observations.")
        doc.add_paragraph()
        
        # Conclusion
        doc.add_heading('3. CONCLUSION', level=2)
        
        # Get rainfall comparison without duplication
        rainfall_text = data.get('rainfall_comparison', '')
        if "This independent evidence" in rainfall_text:
            rainfall_text = rainfall_text.split("This independent evidence")[0].strip()
        
        deviation = data.get('deviation_score', 0)
        if data.get('status') == 'anomaly_detected':
            severity = 'severe' if deviation < -2 else 'significant'
            conclusion = f"""The analysis shows a significant deviation from normal vegetation patterns. 
The NDVI deviation score of {deviation:.2f} standard deviations below the historical baseline 
indicates {severity} vegetation stress.
{rainfall_text}
This independent evidence supports the claim of crop damage and warrants a re-evaluation of the insurance assessment."""
        else:
            conclusion = f"""The analysis shows vegetation health within normal range for this season.
The NDVI deviation score of {deviation:.2f} standard deviations from the historical baseline
indicates normal vegetation conditions.
{rainfall_text}
This independent evidence suggests the insurance assessment should be reviewed in light of normal vegetation conditions."""
        
        doc.add_paragraph(conclusion)
        doc.add_paragraph()
        
        # Request
        doc.add_heading('4. REQUEST', level=2)
        doc.add_paragraph("""In light of the evidence presented above, I respectfully request the following:

1. A formal review of the crop insurance claim assessment for the plot described above
2. Consideration of the independent satellite and weather data as evidence
3. A fair and transparent evaluation based on all available evidence

I am available to provide any additional information or clarification as needed.""")
        doc.add_paragraph()
        
        # Signature
        doc.add_paragraph('Sincerely,')
        doc.add_paragraph()
        doc.add_paragraph('[Farmer Name]')
        doc.add_paragraph('[Contact Number]')
        doc.add_paragraph('[Email Address]')
        doc.add_paragraph('[Village], [District], [State]')
        
        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
    
    def generate_html(self, data: dict) -> str:
        """Generate an HTML version of the appeal document"""
        deviation = data.get('deviation_score', 0)
        if deviation < -2:
            status_text = 'Severe'
        elif deviation < -1.5:
            status_text = 'Significant'
        else:
            status_text = 'Normal'
        
        rainfall_text = data.get('rainfall_comparison', '')
        if "This independent evidence" in rainfall_text:
            rainfall_text = rainfall_text.split("This independent evidence")[0].strip()
        
        plot_name = data.get('plot_name', 'Plot')
        latitude = data.get('latitude', 'N/A')
        longitude = data.get('longitude', 'N/A')
        season = data.get('season', 'Kharif 2017')
        damage_period = data.get('damage_period', 'damage period')
        image_count = data.get('image_count', 0)
        cloud_cover_avg = data.get('cloud_cover_avg', 0)
        rainfall_total = data.get('rainfall_total', 0)
        rainfall_days = data.get('rainfall_days', 0)
        rainfall_comparison = data.get('rainfall_comparison', 'No data')
        weather_source = data.get('weather_source', 'ERA5-Land')
        status = data.get('status', 'normal')
        
        if status == 'anomaly_detected':
            severity = 'severe' if deviation < -2 else 'significant'
            conclusion = f"""The analysis shows a significant deviation from normal vegetation patterns. 
            The NDVI deviation score of {deviation:.2f} standard deviations below the historical baseline 
            indicates {severity} vegetation stress.
            {rainfall_text}
            This independent evidence supports the claim of crop damage and warrants a re-evaluation of the insurance assessment."""
        else:
            conclusion = f"""The analysis shows vegetation health within normal range for this season.
            The NDVI deviation score of {deviation:.2f} standard deviations from the historical baseline
            indicates normal vegetation conditions.
            {rainfall_text}
            This independent evidence suggests the insurance assessment should be reviewed in light of normal vegetation conditions."""
        
        conclusion = conclusion.replace('\n', '<br>')
        
        html = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <title>Appeal Document - FasalPramaan</title>
            <style>
                body {{ font-family: Arial, sans-serif; max-width: 800px; margin: 40px auto; padding: 20px; line-height: 1.6; color: #1a1a2e; }}
                h1 {{ color: #1a472a; text-align: center; font-size: 24px; }}
                h2 {{ color: #2d6a4f; font-size: 18px; margin-top: 20px; }}
                h3 {{ color: #1a472a; font-size: 14px; margin-top: 15px; }}
                .header-line {{ border-top: 2px solid #2d6a4f; margin: 10px 0 20px; }}
                .date {{ text-align: right; font-size: 14px; }}
                .table {{ width: 100%; border-collapse: collapse; margin: 15px 0; }}
                .table th {{ background: #1a472a; color: white; padding: 8px; text-align: center; }}
                .table td {{ padding: 8px; border: 1px solid #ddd; text-align: center; }}
                .table tr:nth-child(even) {{ background: #f8fafc; }}
                .signature {{ margin-top: 40px; }}
                .footer {{ text-align: center; font-size: 12px; color: #94a3b8; margin-top: 30px; border-top: 1px solid #e2e8f0; padding-top: 15px; }}
                ul {{ list-style: none; padding-left: 0; }}
                ul li::before {{ content: "• "; color: #2d6a4f; font-weight: bold; }}
            </style>
        </head>
        <body>
            <h1>🌾 FASALPRAMAAN</h1>
            <p style="text-align:center; color:#64748b;">Independent Crop Insurance Verification</p>
            <div class="header-line"></div>
            
            <p class="date"><strong>Date:</strong> {datetime.now().strftime('%B %d, %Y')}</p>
            
            <h2>To:</h2>
            <p>The District Grievance Redressal Committee<br>
            [District Name], [State]</p>
            
            <h2>Subject: Appeal against crop insurance assessment for {plot_name}</h2>
            
            <p>Respected Members of the Committee,</p>
            
            <p>I am writing to formally appeal the crop insurance assessment for my plot located at 
            latitude {latitude}, longitude {longitude} for the {season} season. 
            Based on independent analysis using satellite imagery and weather data, I am submitting the following evidence 
            for your review.</p>
            
            <h2>1. SATELLITE-BASED VEGETATION ANALYSIS</h2>
            <p>Using Sentinel-2 satellite imagery accessed through Google Earth Engine, 
            the Normalized Difference Vegetation Index (NDVI) was calculated for the period of {damage_period}. 
            NDVI is a widely accepted, peer-reviewed measure of vegetation health.</p>
            
            <table class="table">
                <tr>
                    <th>Metric</th>
                    <th>Value</th>
                    <th>Status</th>
                </tr>
                <tr>
                    <td>NDVI Deviation Score</td>
                    <td>{deviation:.2f} σ</td>
                    <td>{status_text}</td>
                </tr>
                <tr>
                    <td>Cloud-free Images Used</td>
                    <td>{image_count}</td>
                    <td>✓</td>
                </tr>
                <tr>
                    <td>Average Cloud Cover</td>
                    <td>{cloud_cover_avg:.1f}%</td>
                    <td>✓</td>
                </tr>
            </table>
            
            <h2>2. WEATHER DATA ANALYSIS</h2>
            <p>Weather data was obtained from the Open-Meteo API using ERA5-Land reanalysis data.</p>
            <ul>
                <li><strong>Total Rainfall:</strong> {rainfall_total:.1f} mm</li>
                <li><strong>Rainy Days:</strong> {rainfall_days} days</li>
                <li><strong>Comparison:</strong> {rainfall_comparison}</li>
                <li><strong>Data Source:</strong> {weather_source}</li>
            </ul>
            <p>This weather analysis provides independent verification of the satellite observations.</p>
            
            <h2>3. CONCLUSION</h2>
            <p>{conclusion}</p>
            
            <h2>4. REQUEST</h2>
            <p>In light of the evidence presented above, I respectfully request the following:</p>
            <ol>
                <li>A formal review of the crop insurance claim assessment for the plot described above</li>
                <li>Consideration of the independent satellite and weather data as evidence</li>
                <li>A fair and transparent evaluation based on all available evidence</li>
            </ol>
            <p>I am available to provide any additional information or clarification as needed.</p>
            
            <div class="signature">
                <p>Sincerely,</p>
                <p><strong>[Farmer Name]</strong><br>
                [Contact Number]<br>
                [Email Address]<br>
                [Village], [District], [State]</p>
            </div>
            
            <div class="footer">
                Generated by FasalPramaan • Independent Crop Insurance Verification Tool
            </div>
        </body>
        </html>
        """
        return html